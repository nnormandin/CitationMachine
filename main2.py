from copy import deepcopy
import docx

print('enter document name:')
name = input()

# ingest document
doc = docx.Document(name + '.docx')
temp = docx.Document()

def transfer_paragraph(p, doc):
	doc._body._body._insert_p(p._p)

def transfer_run(r, new_par):
	run = new_par.add_run(r.text)
	run.bold = r.bold
	run.italic = r.italic
	run.underline = r.underline
	run.font.name = r.font.name
	run.font.size = r.font.size

for p in doc1.paragraphs:
	sc = False
	for r in p.runs:
		if '{r*}' in r.text:
			print('found')
			sc = True
	if not sc:
		transfer_paragraph(p, doc2)
	else:
		doc2.add_paragraph()


print('enter clone name:')
clone_name = input()
doc2.save(clone_name + '.docx')

