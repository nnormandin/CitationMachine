import os
import time
import docx
import re
import copy
from selenium import webdriver
from selenium.webdriver.common.keys import Keys

# import wait modules
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC

sc = re.compile('\{.+\}')
yr = re.compile('\(.+\)')


class Citation(object):
''' define citation object '''

    def __init__(self, name, caseid, page, year):
        self.name = name
        self.caseid = caseid
        self.page = page
        self.year = year


def insert_citation(r, c):
''' insert a properly formatted citation in text '''

    r.text = c.name
    r2 = copy.deepcopy(r._r)
    r.underline = True
    r2.text = ", {}, {} ({})".format(c.caseid, c.page, c.year)
    r._r.addnext(r2)


def scan_refs(doc):
''' find all of the references located in a document '''

    references = []

    for p in doc.paragraphs:
        for r in p.runs:
            cite = sc.search(r.text)
            if cite:
                references.append(cite.group(0)[1:-1])
    print('located {} citations'.format(len(references)))

    return(references)


def insert_citations(doc, citation_data):
''' insert each citation in document '''

    for p in doc.paragraphs:
        for r in p.runs:
            cite = sc.search(r.text)
            if cite:
                cite = cite.group(0)[1:-1].split('|')
                idx = cite[0].strip()
                pg = cite[1].strip()
                try:
                    c = [x for x in citation_data if x.caseid == idx][0]
                except:
                    print('citation not found for {}'.format(idx))
                    continue
                insert_citation(r, c)

    print("citation insertions complete")


def wait_load(browser, timeout=20, xp="//input[@id='q']"):
''' wait for browser to load search bar '''

    wait = WebDriverWait(browser, timeout)
    element = wait.until(EC.element_to_be_clickable((By.xpath, xp)))


def search_case(w, text):
''' search for the case by caseid '''

    searchbar = w.find_element_by_xpath("//input[@id='q']")
    searchbar.send_keys(text)
    searchbar.send_keys(Keys.ENTER)
    wait_load(w)


def title_to_case(title):
''' turn uppercase title into title-case with lower-case v.'''

    out = []
    for i in title.split():
        if i == "V":
            out.append('v.')
        else:
            out.append(i.title())
    return(' '.join(out))


def parse_case(w, caseid, page):
''' search for and parse a case, returning Citation object '''

    # search for case
    search_case(w, caseid)
    wait_load(w)

    # find title and parse
    xpt = "//h1[@class='title-and-tools-shell']/span"
    title = w.find_element_by_xpath(xpt).text
    title = title_to_case(title)

    # find header data
    xph = "//ct-document-header[@class='ng-isolate-scope']"
    header = w.find_element_by_xpath(xph)

    court = header.find_element_by_xpath("//li[0]").text

    dategroup = header.find_element_by_xpath("//li[2]").text
    year = yr.search(dategroup).group(0)[1:-1]

    print("located {}: {}".format(caseid, title))

    return(Citation(title, caseid, page, year))


def open_document():
''' ask for document name and return opened docx file '''

    print('enter document name')
    docname = input()
    if ".docx" in docname:
        out = docx.Document(docname)
    else:
        out = docx.Document(docname + 'docx')

    print('{} loaded'.format(docname))

    return(out)


def start_webdriver():
''' initialize firefox webdriver '''

    w = webdriver.Firefox()
    print('loading webdriver')
    w.get('https://casetext.com/')
    wait_load(w)
    print('Firefox loaded')
    return(w)


def get_citations(w, refs):
''' loop through citations '''

    out = []

    for i in refs:
        j = i.split('|')
        try:
            out.append(parse_case(w, j[0].strip(), j[1].strip()))
        except:
            print('error locating {} !!!'.format(j[0]))

    return(out)

def save_document(doc):
''' request document save name, default to OUTPUT.docx'''

    print('please enter a document name')
    sname = input()

    if sname:
        if '.docx' in sname:
            doc.save(sname)
        else:
            doc.save(sname + '.docx')
    else:
        doc.save('OUTPUT.docx')



def main():

    # open the document and start the webdriver
    doc = open_document()
    w = start_webdriver()

    # scan document and scrape citation data
    cdata = get_citations(w, scan_refs(doc))

    # insert citations in document
    insert_citations(doc, cdata)

    # ask for save name
    save_document(doc)

if __name__ == "__main__":
    
    main()